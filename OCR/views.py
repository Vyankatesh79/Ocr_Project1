from django.shortcuts import render,redirect
from google.cloud import storage, vision_v1
import json
import re
import os

def base(request):

    return render(request, 'base.html')

os.environ['GOOGLE_APPLICATION_CREDENTIALS'] = r"D:/files/python_file/Django/OCR_New_Google_Vision_Project/office_GV_key/NewServiceAccountToken.json"

import os
import threading
import pythoncom
from win32com.client import Dispatch
from django.conf import settings
from django.shortcuts import render, redirect
from django.http import HttpResponse , JsonResponse , HttpResponseNotFound
from google.cloud import storage, vision_v1
import json
import re
from io import BytesIO
from docx import Document
from reportlab.lib.pagesizes import letter
from reportlab.pdfgen import canvas
from django.conf import settings
from django.http import JsonResponse

from django.http import JsonResponse, FileResponse 
from django.urls import reverse
import shutil

def Get_file(request):
    if request.method == 'POST':
        files = request.FILES.getlist('files')
        print("This is get file name: ", files)
        doc_files = []

        # Upload each file to Google Cloud Storage and process with Google Vision API
        bucket_name = 'ocr_demo1'
        storage_client = storage.Client()
        vision_client = vision_v1.ImageAnnotatorClient()

        for file in files:
            # Upload the file to Google Cloud Storage
            bucket = storage_client.bucket(bucket_name)
            blob = bucket.blob(file.name)
            blob.upload_from_file(file)

            # Process the image with Google Vision API
            mime_type = "application/pdf"
            feature = vision_v1.Feature(type=vision_v1.Feature.Type.DOCUMENT_TEXT_DETECTION)

            gcs_source_uri = f'gs://{bucket_name}/{file.name}'
            gcs_source = vision_v1.GcsSource(uri=gcs_source_uri)
            input_config = vision_v1.InputConfig(gcs_source=gcs_source, mime_type=mime_type)

            gcs_destination_uri = f'gs://{bucket_name}/pdf_result/{file.name}'
            gcs_destination = vision_v1.GcsDestination(uri=gcs_destination_uri)
            output_config = vision_v1.OutputConfig(gcs_destination=gcs_destination, batch_size=1)

            async_request = vision_v1.AsyncAnnotateFileRequest(
                features=[feature], input_config=input_config, output_config=output_config)

            operation = vision_client.async_batch_annotate_files(requests=[async_request])

            print("Waiting for operation to complete...")
            operation.result(timeout=420)

            # Fetch the results from Google Cloud Storage
            match = re.match(r'gs://([^/]+)/(.+)', gcs_destination_uri)
            bucket_name = match.group(1)
            prefix = match.group(2)
            bucket = storage_client.bucket(bucket_name)

            blob_list = list(bucket.list_blobs(prefix=prefix))
            all_extracted_text = ""

            for output in blob_list:
                json_string = output.download_as_string()

                # Parse the JSON response
                response_dict = json.loads(json_string)
                print("This is response: ", response_dict)
                extracted_text = response_dict["responses"][0]["fullTextAnnotation"]["text"]
                all_extracted_text += extracted_text + "\n"

            # Save the extracted text to a single DOC file directly
            media_root = settings.MEDIA_ROOT
            if not os.path.exists(media_root):
                os.makedirs(media_root)

            # Create a unique filename for each DOC file based on the original file name
            base_name = os.path.splitext(file.name)[0]

            doc_file_path = os.path.join(media_root, f'{base_name}.docx')
            
            document = Document()
            document.add_heading('OCR Extracted Text', 0)
            document.add_paragraph(all_extracted_text)
            document.save(doc_file_path)
            
            doc_files.append(doc_file_path)

        # Render the template with a success message and the paths to the DOC files
        return render(request, 'get_file.html', {'success': True, 'doc_files': doc_files})

    return render(request, 'get_file.html')

def convert_doc_to_pdf(doc_file_path, temp_pdf_path):
    try:
        pythoncom.CoInitialize()
        word = Dispatch('Word.Application')
        doc = word.Documents.Open(doc_file_path)
        doc.SaveAs(temp_pdf_path, FileFormat=17)  # 17 is the code for PDF format
        doc.Close()
        word.Quit()
    except Exception as e:
        print(f"Error in conversion: {e}")
    finally:
        pythoncom.CoUninitialize()

def download_file(request):
    media_folder_path = settings.MEDIA_ROOT
    doc_files = [f for f in os.listdir(media_folder_path) if f.endswith('.doc') or f.endswith('.docx')]
    print("DOC files in media folder:", doc_files)
    download = request.POST.get('download')
    print("This is download button request:", download)

    pdf_files = []
    errors = []

    for doc_file_name in doc_files:
        full_doc_file_path = os.path.join(media_folder_path, doc_file_name)
        print(f"Full Doc File Path: {full_doc_file_path}")

        if os.path.exists(full_doc_file_path):
            base_name = os.path.splitext(os.path.basename(doc_file_name))[0]
            temp_pdf_path = os.path.join(settings.MEDIA_ROOT, f'{base_name}.pdf')
            print(f"Temp PDF Path: {temp_pdf_path}")

            convert_doc_to_pdf(full_doc_file_path, temp_pdf_path)

            if os.path.exists(temp_pdf_path):
                pdf_files.append(temp_pdf_path)
                os.remove(full_doc_file_path)  # Delete DOC file after conversion
                print(f"Converted and deleted DOC file: {full_doc_file_path}")
            else:
                errors.append(f"Error converting DOC file: {full_doc_file_path}")
        else:
            errors.append(f"File does not exist: {full_doc_file_path}")

    if errors:
        return JsonResponse({'success': False, 'errors': errors})

    download_urls = [request.build_absolute_uri(reverse('serve_pdf_file', args=[os.path.basename(pdf_file_path)])) for pdf_file_path in pdf_files]
    print("Download URLs:", download_urls)

    return JsonResponse({'success': True, 'download_urls': download_urls})

def serve_pdf_file(request, filename):
    file_path = os.path.join(settings.MEDIA_ROOT, filename)
    if os.path.exists(file_path):
        response = FileResponse(open(file_path, 'rb'), content_type='application/pdf')
        response['Content-Disposition'] = f'attachment; filename="{filename}"'

        # Add a callback to delete the file after the response is completed
        def delete_file_callback(response):
            try:
                os.remove(file_path)
                print(f"File {filename} has been deleted successfully.")
            except OSError as e:
                print(f"Error: {file_path} : {e.strerror}")

        # Add the callback to the response
        response.close = lambda *args, **kwargs: (
            FileResponse.close(response, *args, **kwargs),
            delete_file_callback(response),
        )

        return response
    else:
        return HttpResponseNotFound('File not found')